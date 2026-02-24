import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import json
import re
import PyPDF2

# --- 1. SETTINGS & STYLING ---
st.set_page_config(page_title="Generator Soal SMP Muhammadiyah 1 Weleri", layout="wide")

st.markdown("""
    <style>
    .stApp { background-color: #f8fafc; }
    div.stButton > button {
        background-color: #2563eb; color: white; border-radius: 0.5rem;
        padding: 0.6rem 1.2rem; border: none; font-weight: 600; transition: all 0.2s; width: 100%;
    }
    div.stButton > button:hover {
        background-color: #1d4ed8; transform: translateY(-1px);
        box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
    }
    section[data-testid="stSidebar"] { background-color: #ffffff; border-right: 1px solid #e2e8f0; }
    h1, h2, h3 { color: #1e3a8a !important; font-family: 'Inter', sans-serif; }
    .special-preview {
        padding: 15px; background-color: #ffffff; border-radius: 8px;
        border: 1px solid #e2e8f0; line-height: 1.8;
    }
    .image-placeholder {
        background-color: #fef3c7; border: 2px dashed #d97706;
        padding: 10px; margin: 10px 0; border-radius: 5px; font-style: italic; color: #92400e;
    }
    .arabic-text { font-family: 'Traditional Arabic', serif; direction: rtl; text-align: right; font-size: 24px; }
    .jawa-text { font-family: 'Tuladha Jejeg', 'Noto Sans Javanese', serif; font-size: 22px; text-align: left; }
    </style>
    """, unsafe_allow_html=True)

# --- FUNGSI AMBIL API KEY ---
def get_api_key():
    if "GEMINI_API_KEY" in st.secrets:
        return st.secrets["GEMINI_API_KEY"]
    return ""

# --- UTILS FONT & RTL ---
def set_rtl(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_font(run, size=11, bold=False, is_arabic=False, is_javanese=False):
    if is_arabic:
        run.font.name = 'Traditional Arabic'
        for tag in [qn('w:ascii'), qn('w:hAnsi'), qn('w:cs')]:
            run._element.rPr.rFonts.set(tag, 'Traditional Arabic')
        run.font.size = Pt(size + 5)
    elif is_javanese:
        run.font.name = 'Tuladha Jejeg'
        for tag in [qn('w:ascii'), qn('w:hAnsi'), qn('w:eastAsia')]:
            run._element.rPr.rFonts.set(tag, 'Tuladha Jejeg')
        run.font.size = Pt(size + 4)
    else:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(size)
    run.bold = bold

def clean_json_output(text):
    try:
        text = re.sub(r'```json\s*|\s*```', '', text)
        start = text.find('{')
        end = text.rfind('}') + 1
        return text[start:end] if start != -1 else text
    except: return text

def clean_option_text(opt):
    return re.sub(r'^[A-Ea-eأ-د1-5]\.?\s*', '', str(opt)).strip()

# --- 2. DOKUMEN GENERATORS ---
def create_header(doc, info, title_suffix=""):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MAJELIS PENDIDIKAN DASAR MENENGAH DAN NON FORMAL\n"); set_font(r, 11, True)
    r = p.add_run("PIMPINAN CABANG MUHAMMADIYAH WELERI\n"); set_font(r, 12, True)
    r = p.add_run(f"{info['sekolah']}\n"); set_font(r, 14, True)
    r = p.add_run(f"{info['jenis_asesmen'].upper()} {title_suffix}\n"); set_font(r, 12, True)
    r = p.add_run(f"TAHUN PELAJARAN {info['tahun']}\n"); set_font(r, 11, True)
    doc.add_paragraph("_" * 75)
    
    tbl = doc.add_table(3, 2)
    rows = [(f"MATA PELAJARAN : {info['mapel']}", f"WAKTU : 90 Menit"), 
            (f"HARI/TANGGAL : .................", f"KELAS : {info['kelas']} / {info['semester']}"), 
            (f"GURU PENGAMPU : {info['guru']}", "")]
    for i, (left, right) in enumerate(rows):
        set_font(tbl.rows[i].cells[0].paragraphs[0].add_run(left), 10)
        set_font(tbl.rows[i].cells[1].paragraphs[0].add_run(right), 10)
    doc.add_paragraph()

def generate_naskah(data_list, info):
    doc = Document(); create_header(doc, info)
    is_arab = "arab" in info['mapel'].lower()
    is_jawa = "jawa" in info['mapel'].lower()
    grouped = {}
    for q in data_list:
        t = q.get('tipe', 'Soal'); grouped.setdefault(t, []).append(q)
    no = 1
    for tipe, quests in grouped.items():
        doc.add_paragraph().add_run(f"\n{tipe.upper()}").bold = True
        for q in quests:
            p_soal = doc.add_paragraph()
            if is_arab: set_rtl(p_soal)
            run_no = p_soal.add_run(f"{no}. "); set_font(run_no, 11)
            
            # CEK APAKAH ADA INSTRUKSI GAMBAR
            soal_text = q.get('soal', '')
            if "[Gambar:" in soal_text:
                parts = soal_text.split("[Gambar:")
                run_text = p_soal.add_run(parts[0]); set_font(run_text, 12, is_arabic=is_arab, is_javanese=is_jawa)
                
                # Tambah Kotak Kosong untuk Gambar di Word
                p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_img = p_img.add_run(f"\n[ TEMPAT GAMBAR: {parts[1].split(']')[0]} ]\n")
                set_font(r_img, 10, bold=True)
                
                if len(parts[1].split(']')) > 1:
                    run_text2 = p_soal.add_run(parts[1].split(']')[1]); set_font(run_text2, 12, is_arabic=is_arab, is_javanese=is_jawa)
            else:
                run_text = p_soal.add_run(soal_text); set_font(run_text, 12, is_arabic=is_arab, is_javanese=is_jawa)

            if "Pilihan Ganda" in tipe:
                opsi = q.get('opsi', [])
                labels = ['أ', 'ب', 'ج', 'د'] if is_arab else ['A', 'B', 'C', 'D']
                for i, o in enumerate(opsi[:4]):
                    p_opt = doc.add_paragraph()
                    if is_arab: set_rtl(p_opt)
                    run_opt = p_opt.add_run(f"    {labels[i]}. {clean_option_text(o)}"); set_font(run_opt, 11, is_arabic=is_arab, is_javanese=is_jawa)
            elif "Isian" in tipe or "Uraian" in tipe:
                doc.add_paragraph("    ......................................................................................................")
            no += 1
    return doc

# --- FUNGSI LAIN (KUNCI, KISI, KARTU) TETAP SAMA ---
def generate_kunci_pedoman(data_list, info):
    doc = Document(); create_header(doc, info, "- KUNCI JAWABAN & PEDOMAN")
    is_arab = "arab" in info['mapel'].lower()
    is_jawa = "jawa" in info['mapel'].lower()
    table = doc.add_table(1, 4); table.style = 'Table Grid'
    for i, h in enumerate(["No", "Tipe", "Kunci/Pedoman", "Skor"]): table.rows[0].cells[i].text = h
    for i, q in enumerate(data_list):
        row = table.add_row().cells
        row[0].text = str(i+1); row[1].text = q.get('tipe', '-')
        p = row[2].paragraphs[0]
        if is_arab: set_rtl(p)
        txt = f"Kunci: {q.get('kunci', '')}\nPedoman: {q.get('pedoman', '')}"
        run = p.add_run(txt); set_font(run, 10, is_arabic=is_arab, is_javanese=is_jawa)
        row[3].text = str(round(q.get('skor', 0), 1))
    return doc

def generate_kisi_kisi(data_list, info):
    doc = Document(); doc.add_heading(f"KISI-KISI SOAL {info['jenis_asesmen']}", 1)
    table = doc.add_table(1, 6); table.style = 'Table Grid'
    for i, h in enumerate(["No", "TP/KD", "Indikator", "Level", "Bentuk", "No Soal"]): table.rows[0].cells[i].text = h
    for i, q in enumerate(data_list):
        row = table.add_row().cells
        for idx, val in enumerate([str(i+1), q.get('tp','-'), q.get('indikator','-'), q.get('level','L2'), q.get('tipe','-'), str(i+1)]): row[idx].text = val
    return doc

def generate_kartu(data_list, info):
    doc = Document(); is_arab = "arab" in info['mapel'].lower(); is_jawa = "jawa" in info['mapel'].lower()
    for i, q in enumerate(data_list):
        doc.add_heading(f"KARTU SOAL - {info['guru']}", 1)
        tbl = doc.add_table(5, 2); tbl.style = 'Table Grid'
        cells = [("Nomor", str(i+1)), ("Indikator", q.get('indikator','-')), ("Butir Soal", q.get('soal','-')), ("Kunci", q.get('kunci','-')), ("Skor", str(round(q.get('skor',0),1)))]
        for idx, (l, v) in enumerate(cells):
            tbl.cell(idx, 0).text = l
            p = tbl.cell(idx, 1).paragraphs[0]
            if is_arab and idx in [2,3]: set_rtl(p)
            run = p.add_run(str(v)); set_font(run, 11, is_arabic=(is_arab and idx in [2,3]), is_javanese=(is_jawa and idx in [2,3]))
        doc.add_page_break()
    return doc

# --- 3. UI UTAMA ---
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/a/af/Muhammadiyah_Logo.svg/1200px-Muhammadiyah_Logo.svg.png", width=80)
    st.header("⚙️ Konfigurasi")
    st.divider()
    sekolah = st.text_input("Sekolah", "SMP MUHAMMADIYAH 1 WELERI")
    guru = st.text_input("Guru", "Ary Prabowo")
    mapel = st.text_input("Mapel", "Bahasa Jawa")
    kelas = st.text_input("Kelas", "IX")
    semester = st.selectbox("Semester", ["Gasal", "Genap"])
    tahun = st.text_input("Tahun", "2025/2026")

st.markdown("<h1 style='text-align: center;'>📝 Generator Administrasi Soal AI</h1>", unsafe_allow_html=True)

with st.container():
    st.markdown("<div style='background-color: white; padding: 2rem; border-radius: 1rem; border: 1px solid #e2e8f0;'>", unsafe_allow_html=True)
    st.subheader("📖 Input Materi")
    c_m1, c_m2 = st.columns(2)
    with c_m1: materi_manual = st.text_area("Teks Materi", height=200)
    with c_m2: uploaded_file = st.file_uploader("Upload PDF", type=['pdf'])
    st.markdown("</div>", unsafe_allow_html=True)

st.write("")
col_s1, col_s2 = st.columns([2, 1])
with col_s1:
    jenis_asesmen = st.selectbox("Peruntukan Soal", ["Asesmen Formatif", "Asesmen Sumatif Lingkup Materi", "Asesmen Sumatif Tengah Semester", "Asesmen Sumatif Akhir Semester"])
    # TAMBAHKAN OPSI SOAL BERGAMBAR
    bentuk_soal = st.multiselect("Bentuk Soal", ["Pilihan Ganda", "Pilihan Ganda (Bergambar)", "Pilihan Ganda Kompleks", "Benar / Salah", "Isian Singkat", "Uraian"], default=["Pilihan Ganda", "Uraian"])
with col_s2:
    conf = {b: st.number_input(f"Jumlah {b}", 0, 30, 0 if "Bergambar" in b else 5) for b in bentuk_soal}

if st.button("🚀 GENERATE SEKARANG"):
    api_key = get_api_key()
    if not api_key: st.error("API Key tidak ditemukan!"); st.stop()
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        materi_full = materi_manual + " "
        if uploaded_file:
            reader = PyPDF2.PdfReader(uploaded_file)
            materi_full += " ".join([p.extract_text() for p in reader.pages])

        is_arab = "arab" in mapel.lower()
        is_jawa = "jawa" in mapel.lower()
        instr = ""
        if is_arab: instr = "GUNAKAN BAHASA ARAB BERHARAKAT."
        if is_jawa: instr = "GUNAKAN BAHASA JAWA DAN SERTAKAN AKSARA JAWA (UNICODE)."
        
        # PROMPT KHUSUS GAMBAR
        prompt_gambar = "Untuk 'Pilihan Ganda (Bergambar)', buat soal yang mengacu pada stimulus visual. Tuliskan instruksi gambar di awal soal dengan format [Gambar: Deskripsi Gambar]."

        prompt = f"""Buat soal {mapel} untuk {jenis_asesmen}. Materi: {materi_full[:7000]}. 
        Bentuk soal: {json.dumps(conf)}. {instr} {prompt_gambar}
        Aturan: Total skor 100. Output JSON MURNI: {{ "soal_list": [ {{ "tipe": "", "soal": "", "opsi": [], "kunci": "", "pedoman": "", "indikator": "", "tp": "", "skor": 0, "level": "" }} ] }}"""

        with st.spinner("AI sedang menyusun soal bergambar..."):
            res = model.generate_content(prompt)
            data = json.loads(clean_json_output(res.text))
            soal_list = data.get('soal_list', [])
            
            info = {'sekolah': sekolah, 'guru': guru, 'mapel': mapel, 'kelas': kelas, 'semester': semester, 'tahun': tahun, 'jenis_asesmen': jenis_asesmen}
            st.session_state.preview_data = soal_list
            st.session_state.mapel_jawa = is_jawa
            st.session_state.mapel_arab = is_arab
            st.session_state.files = {'n': generate_naskah(soal_list, info), 'k': generate_kisi_kisi(soal_list, info), 's': generate_kartu(soal_list, info), 'kj': generate_kunci_pedoman(soal_list, info)}
            st.success("🎉 Berhasil Dibuat!")
    except Exception as e:
        st.error(f"Kesalahan: {e}")

if 'files' in st.session_state:
    st.divider()
    c1, c2, c3, c4 = st.columns(4)
    def to_io(doc):
        io = BytesIO(); doc.save(io); return io.getvalue()
    with c1: st.download_button("📄 Naskah", to_io(st.session_state.files['n']), "Naskah.docx")
    with c2: st.download_button("🔑 Kunci", to_io(st.session_state.files['kj']), "Kunci.docx")
    with c3: st.download_button("📋 Kisi", to_io(st.session_state.files['k']), "Kisi.docx")
    with c4: st.download_button("🗃️ Kartu", to_io(st.session_state.files['s']), "Kartu.docx")

    st.subheader("👁️ Preview Soal")
    for i, q in enumerate(st.session_state.preview_data):
        with st.expander(f"Soal {i+1} - {q.get('tipe')}"):
            s_class = "arabic-text" if st.session_state.mapel_arab else ("jawa-text" if st.session_state.mapel_jawa else "")
            
            soal_display = q['soal']
            if "[Gambar:" in soal_display:
                parts = soal_display.split("[Gambar:")
                st.write(parts[0])
                st.markdown(f"<div class='image-placeholder'>🖼️ <b>Instruksi Gambar:</b> {parts[1].split(']')[0]}</div>", unsafe_allow_html=True)
                if len(parts[1].split(']')) > 1: st.markdown(f"<div class='{s_class}'>{parts[1].split(']')[1]}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f<div class='{s_class}'>{soal_display}</div>", unsafe_allow_html=True)
            
            if q.get('opsi'):
                for idx, opt in enumerate(q.get('opsi')):
                    st.markdown(f"<div class='{s_class}'>{['A','B','C','D'][idx]}. {clean_option_text(opt)}</div>", unsafe_allow_html=True)
            st.info(f"Kunci: {q.get('kunci')}")
